# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_LINE_SPACING
import os

# 读取源代码文件
def read_file(path):
    try:
        with open(path, 'r', encoding='utf-8') as f:
            return f.readlines()
    except:
        return []

# 创建Word文档
def create_doc(title, files_content, output_path, lines_per_page=50, total_pages=30):
    doc = Document()

    # 设置默认字体和大小（宋体小五号 = 9磅）
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(9)

    # 设置页边距（上下左右各2.5cm）
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.984)  # 2.5cm
        section.bottom_margin = Inches(0.984)  # 2.5cm
        section.left_margin = Inches(0.984)  # 2.5cm
        section.right_margin = Inches(0.984)  # 2.5cm

    # 添加标题
    doc.add_heading(title, 0)

    current_page = 1
    lines_on_current_page = 0
    total_lines_needed = lines_per_page * total_pages

    for file_path, lines in files_content:
        if lines:
            # 添加文件分隔注释
            para = doc.add_paragraph()
            run = para.add_run(f'\n// ========== {file_path} ==========')
            run.font.size = Pt(8)
            run.font.name = '宋体'
            run.italic = True
            lines_on_current_page += 1

            for line in lines:
                line = line.rstrip()

                # 每50行换一页
                if lines_on_current_page >= lines_per_page:
                    doc.add_page_break()
                    current_page += 1
                    lines_on_current_page = 0

                    # 添加页眉
                    para = doc.add_paragraph()
                    run = para.add_run(f'软件名称《心伴》V1.0 - 第 {current_page} 页')
                    run.font.size = Pt(8)
                    run.font.name = '宋体'
                    lines_on_current_page += 1

                # 添加代码行
                para = doc.add_paragraph()
                para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                para.paragraph_format.line_spacing = Pt(14)  # 固定值14磅

                # 处理制表符，替换为4个空格
                line = line.replace('\t', '    ')
                run = para.add_run(line)
                run.font.size = Pt(9)
                run.font.name = '宋体'

                lines_on_current_page += 1

    # 保存文档
    doc.save(output_path)
    print(f'文档已生成: {output_path}')
    return output_path

# 前30页源代码 - 核心功能代码
files_front_30 = [
    ('src/app/page.tsx', read_file('src/app/page.tsx')),
    ('src/components/timer/Timer.tsx', read_file('src/components/timer/Timer.tsx')),
    ('src/components/chat/ChatWindow.tsx', read_file('src/components/chat/ChatWindow.tsx')),
    ('src/app/calendar/page.tsx', read_file('src/app/calendar/page.tsx')),
    ('src/app/profile/page.tsx', read_file('src/app/profile/page.tsx')),
    ('src/app/journal/page.tsx', read_file('src/app/journal/page.tsx')),
    ('src/app/tasks/page.tsx', read_file('src/app/tasks/page.tsx')),
]

# 后30页源代码 - 配置和工具代码
files_back_30 = [
    ('src/components/timer/TimerStats.tsx', read_file('src/components/timer/TimerStats.tsx')),
    ('src/components/chat/MessageBubble.tsx', read_file('src/components/chat/MessageBubble.tsx')),
    ('src/components/ui/BottomNav.tsx', read_file('src/components/ui/BottomNav.tsx')),
    ('src/components/ui/button.tsx', read_file('src/components/ui/button.tsx')),
    ('src/components/ui/card.tsx', read_file('src/components/ui/card.tsx')),
    ('src/components/ui/input.tsx', read_file('src/components/ui/input.tsx')),
    ('src/components/ui/dialog.tsx', read_file('src/components/ui/dialog.tsx')),
    ('src/components/ui/drawer.tsx', read_file('src/components/ui/drawer.tsx')),
    ('src/types/index.ts', read_file('src/types/index.ts')),
    ('src/components/radio/WhiteNoisePlayer.tsx', read_file('src/components/radio/WhiteNoisePlayer.tsx')),
    ('src/app/rant/page.tsx', read_file('src/app/rant/page.tsx')),
    ('src/app/layout.tsx', read_file('src/app/layout.tsx')),
    ('src/components/mood/MoodCheckIn.tsx', read_file('src/components/mood/MoodCheckIn.tsx')),
    ('src/app/globals.css', read_file('src/app/globals.css')),
]

# 生成前30页文档
print('正在生成前30页源代码文档...')
print(f'代码文件数量: {len(files_front_30)}')
total_lines_front = sum(len(lines) for _, lines in files_front_30)
print(f'总行数: {total_lines_front}')
create_doc(
    '心伴_源代码前30页',
    files_front_30,
    'xinban_source_front_30pages.docx',
    lines_per_page=50,
    total_pages=30
)

# 生成后30页文档
print('\n正在生成后30页源代码文档...')
print(f'代码文件数量: {len(files_back_30)}')
total_lines_back = sum(len(lines) for _, lines in files_back_30)
print(f'总行数: {total_lines_back}')
create_doc(
    '心伴_源代码后30页',
    files_back_30,
    'xinban_source_back_30pages.docx',
    lines_per_page=50,
    total_pages=30
)

print('\n✅ 所有文档生成完成！')
print('请查看生成的文件：')
print('  - xinban_source_front_30pages.docx')
print('  - xinban_source_back_30pages.docx')
print(f'\n前30页总行数: {total_lines_front}')
print(f'后30页总行数: {total_lines_back}')